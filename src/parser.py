import re
from pathlib import Path
import pypdf
import docx
import openpyxl

def clean_whitespace(text: str) -> str:
    """Replaces multiple whitespaces and newlines with single instances to clean text."""
    text = re.sub(r'[ \t]+', ' ', text)
    text = re.sub(r'\n+', '\n', text)
    return text.strip()

def extract_text_from_txt(filepath: Path) -> str:
    """Reads a text or markdown file, trying different encodings.

    Order matters: utf-8-sig handles UTF-8 with or without a BOM, utf-16 is
    tried (BOM-aware) before latin1, and latin1 is kept LAST as the catch-all.
    latin1 never raises UnicodeDecodeError, so anything placed after it would be
    unreachable and real UTF-16 files would be silently mangled.
    """
    encodings = ['utf-8-sig', 'utf-16', 'cp949', 'euc-kr', 'latin1']
    for encoding in encodings:
        try:
            with open(filepath, 'r', encoding=encoding) as f:
                return f.read()
        except UnicodeDecodeError:
            continue
    raise ValueError(f"Could not decode text file {filepath} with supported encodings.")

# Optional OCR fallback for scanned/image PDFs. The dependencies (PyMuPDF for
# page rendering, rapidocr-onnxruntime for recognition) are heavy and optional;
# they are imported lazily so the core works without them. When absent, scanned
# PDFs are simply detected and skipped (logged), not OCR'd.
_ocr_engine = None
_ocr_unavailable = False

def _get_ocr_engine():
    """Returns a cached RapidOCR engine, or None if rapidocr is not installed."""
    global _ocr_engine, _ocr_unavailable
    if _ocr_engine is not None:
        return _ocr_engine
    if _ocr_unavailable:
        return None
    try:
        from rapidocr_onnxruntime import RapidOCR
        _ocr_engine = RapidOCR()
        return _ocr_engine
    except Exception as e:
        print(f"OCR fallback unavailable (rapidocr-onnxruntime not installed): {e}")
        _ocr_unavailable = True
        return None

def ocr_pdf(filepath: Path) -> str:
    """Renders each PDF page to an image and runs OCR. Returns text or '' if OCR
    is unavailable or fails."""
    try:
        import fitz  # PyMuPDF
    except Exception:
        print("OCR fallback skipped: PyMuPDF (fitz) not installed.")
        return ""
    engine = _get_ocr_engine()
    if engine is None:
        return ""

    texts = []
    try:
        doc = fitz.open(str(filepath))
        try:
            for page in doc:
                pix = page.get_pixmap(dpi=200)
                result, _ = engine(pix.tobytes("png"))
                if result:
                    texts.append(" ".join(line[1] for line in result))
        finally:
            doc.close()
    except Exception as e:
        print(f"OCR failed for {filepath}: {e}")
        return ""
    return "\n".join(texts)

def extract_text_from_pdf(filepath: Path) -> str:
    """Extracts text content from a PDF file, falling back to OCR for scans."""
    texts = []
    try:
        with open(filepath, 'rb') as f:
            reader = pypdf.PdfReader(f)
            for page in reader.pages:
                try:
                    text = page.extract_text()
                    if text:
                        texts.append(text)
                except Exception as e:
                    print(f"Warning: Failed to extract page text in PDF {filepath}: {e}")
                    continue
    except Exception as e:
        print(f"Error opening PDF {filepath}: {e}")

    combined = "\n".join(texts)
    if combined.strip():
        return combined

    # No embedded text layer -- likely a scanned/image PDF. Try OCR if available.
    print(f"No text layer in PDF {filepath}; attempting OCR fallback...")
    ocr_text = ocr_pdf(filepath)
    if ocr_text.strip():
        print(f"OCR extracted {len(ocr_text)} chars from {filepath}.")
        return ocr_text
    print(f"No extractable text from PDF {filepath} (scanned and OCR unavailable, or empty).")
    return combined

def extract_text_from_docx(filepath: Path) -> str:
    """Extracts text content from paragraphs and tables in a Word document."""
    doc = docx.Document(str(filepath))
    texts = []
    
    # Extract paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            texts.append(para.text)
            
    # Extract tables
    for table in doc.tables:
        for row in table.rows:
            row_texts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_texts:
                texts.append(" | ".join(row_texts))
                
    return "\n".join(texts)

def extract_text_from_xlsx(filepath: Path) -> str:
    """Extracts text content from all sheets in an Excel spreadsheet."""
    wb = openpyxl.load_workbook(str(filepath), data_only=True, read_only=True)
    try:
        texts = []
        for sheet in wb.worksheets:
            texts.append(f"Sheet: {sheet.title}")
            for row in sheet.iter_rows(values_only=True):
                row_text = " ".join([str(val).strip() for val in row if val is not None])
                if row_text.strip():
                    texts.append(row_text)
        return "\n".join(texts)
    finally:
        wb.close()

def extract_text(filepath_str: str) -> str:
    """Dispatches text extraction based on file extension."""
    filepath = Path(filepath_str)
    ext = filepath.suffix.lower()
    
    if ext in {'.txt', '.md'}:
        return extract_text_from_txt(filepath)
    elif ext == '.pdf':
        return extract_text_from_pdf(filepath)
    elif ext == '.docx':
        return extract_text_from_docx(filepath)
    elif ext == '.xlsx':
        return extract_text_from_xlsx(filepath)
    else:
        raise ValueError(f"Unsupported file format: {ext}")

def chunk_text(text: str, chunk_size: int = 500, overlap: int = 50) -> list[str]:
    """Splits cleaned text into chunks of specified size and overlap."""
    text = clean_whitespace(text)
    if not text:
        return []

    # Guard against a misconfiguration where overlap >= chunk_size, which would
    # make the step non-positive and loop forever.
    step = chunk_size - overlap
    if step <= 0:
        step = chunk_size

    chunks = []
    start = 0
    text_len = len(text)

    while start < text_len:
        end = start + chunk_size
        chunk = text[start:end]
        chunks.append(chunk)

        if end >= text_len:
            break

        start += step

    return chunks
